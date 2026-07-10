import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def force_table_alignment():
    doc = Document('实验报告.docx')
    t = doc.tables[0]
    for row in t.rows:
        # Clear row height
        row.height = None
        row.height_rule = None
        if row._tr.trPr is not None:
            for child in row._tr.trPr:
                if child.tag.endswith('trHeight'):
                    row._tr.trPr.remove(child)
                    
        for cell in row.cells:
            # Force vertical alignment center in XML
            tcPr = cell._tc.get_or_add_tcPr()
            vAlign = tcPr.find(qn('w:vAlign'))
            if vAlign is None:
                vAlign = OxmlElement('w:vAlign')
                tcPr.append(vAlign)
            vAlign.set(qn('w:val'), 'center')
            
            for p in cell.paragraphs:
                p_fmt = p.paragraph_format
                p_fmt.space_before = 0
                p_fmt.space_after = 0
                p_fmt.line_spacing = 1.0
                
    doc.save('实验报告.docx')
    print("Forced table alignment and removed row heights")

if __name__ == '__main__':
    force_table_alignment()
