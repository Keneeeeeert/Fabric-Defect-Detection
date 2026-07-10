import sys
from docx import Document

def change_name():
    try:
        doc = Document('实验报告_完整版.docx')
        for p in doc.paragraphs:
            if '施哲诚' in p.text:
                p.text = p.text.replace('施哲诚', 'Keneeeeert')
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '施哲诚' in cell.text:
                        cell.text = cell.text.replace('施哲诚', 'Keneeeeert')
        doc.save('实验报告_完整版.docx')
        print("Name replaced in docx")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    change_name()
