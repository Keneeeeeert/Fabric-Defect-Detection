import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_font_yahei(run):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def center_and_font_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in p.runs:
                    set_font_yahei(run)

def polish_doc():
    doc = Document('实验报告.docx')
    
    # Define polishing dictionary (from my generated text to polished high school tone)
    replacements = {
        '本实验采用的模型为：YOLOv8n (常规框) 和 YOLOv8n-obb (旋转框)。': '本实验采用了目前前沿的深度学习模型：YOLOv8n（用于常规目标检测）以及专门针对细长目标改进的 YOLOv8n-obb（旋转框检测模型）。',
        'mAP50 达 0.928': '能够非常准确地识别，mAP50（平均精度）达到了 0.928',
        '由于边界模糊，mAP50 仅 0.444': '由于污渍边缘比较模糊，容易和衣服纹理混淆，目前的 mAP50 为 0.444，还有进步空间',
        '使用 OBB 后 mAP50 达 0.859': '原本极难检测的线头，在使用 OBB 旋转框技术后，mAP50 跃升至 0.859',
        'mAP50: 0.465 (常规) / 0.859 (OBB)': '总体平均精度(mAP50)达到 0.465，线头专项达到 0.859',
        'Hole类高，Dirt类较低': '破洞（Hole）检测精确率较高，污渍（Dirt）因为样本特征不明显而稍低',
        '良好，尤其旋转框大幅提升召回率': '表现良好，尤其是引入旋转框技术后，线头的召回率大幅度提升',
        'Dirt 漏检主要由于数据量不足': '目前的漏检主要集中在污渍（Dirt）类别，主要是因为采集的数据样本还不够多',
        '初始阶段，Colab batch=32': '模型刚开始训练，我们在 Colab 上设置了一次训练 32 张图片（Batch Size = 32）',
        '指标稳步上升': '模型逐渐学会了识别疵点，各项精度指标正在稳步上升',
        '模型收敛，早停触发前最好结果': '模型基本学成（收敛），达到了目前的最好识别效果'
    }

    # Polish and format paragraphs
    for p in doc.paragraphs:
        # Polish
        text = p.text
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
        if text != p.text:
            p.text = text
            
        # Format title/headers
        if '实验报告' in p.text and len(p.text) < 15:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        for run in p.runs:
            set_font_yahei(run)

    # Center and format tables
    for table in doc.tables:
        center_and_font_table(table)
        
    doc.save('实验报告_Final.docx')
    print('Polished document saved to 实验报告_Final.docx')

if __name__ == '__main__':
    polish_doc()
