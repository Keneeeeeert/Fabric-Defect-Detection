import sys
import os
from docx import Document

def revert_name_and_rename():
    try:
        # 1. Open the document
        doc = Document('实验报告_完整版.docx')
        
        # 2. Replace Keneeeeert back to real name
        for p in doc.paragraphs:
            if 'Keneeeeert' in p.text:
                p.text = p.text.replace('Keneeeeert', '施哲诚')
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'Keneeeeert' in cell.text:
                        cell.text = cell.text.replace('Keneeeeert', '施哲诚')
                        
        # 3. Save as 实验报告.docx
        doc.save('实验报告.docx')
        print("Name reverted and saved as 实验报告.docx")
        
        # 4. Remove the _完整版 file if possible
        try:
            os.remove('实验报告_完整版.docx')
            print("Removed 实验报告_完整版.docx")
        except Exception as e:
            print(f"Could not remove _完整版: {e}")
            
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    revert_name_and_rename()
