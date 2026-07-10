import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_run_font(run, ascii_font, eastasia_font, size_pt):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), eastasia_font)

def format_doc():
    doc = Document('实验报告.docx')
    
    title_prefixes = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十', '附录']
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        is_main_title = (text == '实验报告')
        is_heading = any(text.startswith(prefix) for prefix in title_prefixes) and len(text) < 25
        
        # Determine font and size
        if is_main_title:
            ascii_font = 'Times New Roman'
            eastasia_font = '黑体'
            size_pt = 22 # 二号
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif is_heading:
            ascii_font = 'Times New Roman'
            eastasia_font = '黑体'
            size_pt = 16 # 三号
        else:
            ascii_font = 'Times New Roman'
            eastasia_font = '宋体'
            size_pt = 12 # 小四
            
        for run in p.runs:
            set_run_font(run, ascii_font, eastasia_font, size_pt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in p.runs:
                        set_run_font(run, 'Times New Roman', '宋体', 10.5) # 五号
                        
    doc.save('实验报告_Standard.docx')
    print('Standard formatted document saved to 实验报告_Standard.docx')

if __name__ == '__main__':
    format_doc()
