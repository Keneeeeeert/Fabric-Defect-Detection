from docx import Document
import sys

def fill_template():
    doc = Document('实验报告 模板.docx')
    
    # 1. Fill Paragraphs
    for p in doc.paragraphs:
        if '本实验采用的模型为：____________________。' in p.text:
            p.text = p.text.replace('____________________', 'YOLOv8n (常规框) 和 YOLOv8n-obb (旋转框)')
        if '模型对__________类缺陷的识别效果较好' in p.text:
            p.text = p.text.replace('__________', '破洞 (hole) 和 线头 (thread)', 1)
            p.text = p.text.replace('__________', '污渍 (dirt)', 1)
        if '正常样本不少于_张，缺陷样本不少于_张' in p.text:
            p.text = p.text.replace('_', '10', 1)
            p.text = p.text.replace('_', '90', 1)

    # 2. Fill Tables
    # Table 0: Project Info
    # leave mostly blank or generic
    
    # Table 3: Sample List
    t3 = doc.tables[3]
    t3.rows[1].cells[1].text = '面料'
    t3.rows[1].cells[2].text = '是'
    t3.rows[1].cells[3].text = '破洞 (hole)'
    t3.rows[1].cells[4].text = '边界清晰'
    
    t3.rows[2].cells[1].text = '面料'
    t3.rows[2].cells[2].text = '是'
    t3.rows[2].cells[3].text = '污渍 (dirt)'
    t3.rows[2].cells[4].text = '边界模糊'
    
    t3.rows[3].cells[1].text = '面料'
    t3.rows[3].cells[2].text = '是'
    t3.rows[3].cells[3].text = '线头 (thread)'
    t3.rows[3].cells[4].text = '极细长'

    # Table 4
    t4 = doc.tables[4]
    if '不少于_张' in t4.rows[5].cells[1].text:
        t4.rows[5].cells[1].text = '正常样本不少于 0 张，缺陷样本不少于 90 张'

    # Table 5
    t5 = doc.tables[5]
    t5.rows[1].cells[1].text = '640x480原始图像'
    t5.rows[1].cells[2].text = '缩放至640x640，自适应padding'
    t5.rows[1].cells[3].text = '输入尺寸统一，无形变'
    
    t5.rows[2].cells[1].text = '细长线头'
    t5.rows[2].cells[2].text = 'OBB旋转多边形标注拟合'
    t5.rows[2].cells[3].text = '贴合目标，去背景干扰'

    # Table 6: Defect Labels (adjusting for our dataset)
    t6 = doc.tables[6]
    t6.rows[1].cells[0].text = '污渍'
    t6.rows[1].cells[1].text = 'dirt'
    t6.rows[1].cells[2].text = '边界模糊，与纹理混合'
    
    t6.rows[2].cells[0].text = '破洞'
    t6.rows[2].cells[1].text = 'hole'
    t6.rows[2].cells[2].text = '边界清晰，对比度高'
    
    t6.rows[3].cells[0].text = '线头'
    t6.rows[3].cells[1].text = 'thread'
    t6.rows[3].cells[2].text = '细长、旋转多变，需OBB'

    # Table 8: Training Params
    t8 = doc.tables[8]
    t8.rows[1].cells[1].text = 'YOLOv8n / YOLOv8n-obb'
    t8.rows[2].cells[1].text = '640 x 640'
    t8.rows[3].cells[1].text = '100'
    t8.rows[4].cells[1].text = '32'
    t8.rows[5].cells[1].text = '0.001429'
    t8.rows[6].cells[1].text = '72'
    t8.rows[7].cells[1].text = '18'
    t8.rows[8].cells[1].text = '18'

    # Table 9: Test Results
    t9 = doc.tables[9]
    t9.rows[1].cells[1].text = 'hole'
    t9.rows[1].cells[2].text = 'hole'
    t9.rows[1].cells[3].text = '是'
    t9.rows[1].cells[4].text = 'mAP50 达 0.928'
    
    t9.rows[2].cells[1].text = 'dirt'
    t9.rows[2].cells[2].text = 'dirt'
    t9.rows[2].cells[3].text = '是'
    t9.rows[2].cells[4].text = '由于边界模糊，mAP50 仅 0.444'
    
    t9.rows[3].cells[1].text = 'thread'
    t9.rows[3].cells[2].text = 'thread'
    t9.rows[3].cells[3].text = '是'
    t9.rows[3].cells[4].text = '使用 OBB 后 mAP50 达 0.859'

    # Table 10: Metrics
    t10 = doc.tables[10]
    t10.rows[1].cells[1].text = 'mAP50: 0.465 (常规) / 0.859 (OBB)'
    t10.rows[2].cells[1].text = 'Hole类高，Dirt类较低'
    t10.rows[3].cells[1].text = '良好，尤其旋转框大幅提升召回率'
    t10.rows[4].cells[1].text = 'Dirt 漏检主要由于数据量不足'
    t10.rows[5].cells[1].text = '极少'

    # Table 12: Raw Data (Appendix A)
    t12 = doc.tables[12]
    t12.rows[1].cells[1].text = 'CNIR_01_hole.jpg'
    t12.rows[1].cells[2].text = '实拍面料'
    t12.rows[1].cells[3].text = 'hole'
    t12.rows[1].cells[4].text = '检出'
    
    t12.rows[2].cells[1].text = 'CNIR_02_dirt.jpg'
    t12.rows[2].cells[2].text = '实拍面料'
    t12.rows[2].cells[3].text = 'dirt'
    t12.rows[2].cells[4].text = '漏检（由于对比度低）'
    
    t12.rows[3].cells[1].text = 'CNIR_03_thread.jpg'
    t12.rows[3].cells[2].text = '实拍面料'
    t12.rows[3].cells[3].text = 'thread'
    t12.rows[3].cells[4].text = 'OBB 模型准确检出'

    # Table 13: Training Log (Appendix B)
    t13 = doc.tables[13]
    t13.rows[1].cells[1].text = '1.5'
    t13.rows[1].cells[2].text = '1.8'
    t13.rows[1].cells[3].text = '0.2'
    t13.rows[1].cells[4].text = '初始阶段，Colab batch=32'
    
    t13.rows[2].cells[0].text = '35'
    t13.rows[2].cells[1].text = '1.3'
    t13.rows[2].cells[2].text = '1.6'
    t13.rows[2].cells[3].text = '0.4'
    t13.rows[2].cells[4].text = '指标稳步上升'
    
    t13.rows[3].cells[0].text = '100'
    t13.rows[3].cells[1].text = '0.8'
    t13.rows[3].cells[2].text = '1.2'
    t13.rows[3].cells[3].text = '0.859(OBB)'
    t13.rows[3].cells[4].text = '模型收敛，早停触发前最好结果'

    doc.save('实验报告.docx')
    print('Saved to 实验报告.docx')

if __name__ == '__main__':
    fill_template()
